#!/usr/bin/python3
import argparse
import os
import glob
import shutil
import xml.etree.ElementTree as ET
import sys
import string
from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.table import WD_TABLE_ALIGNMENT
import re
import collections

#TODO: make lowest score an argument

#Set up Argparser
parser = argparse.ArgumentParser(description="A script to create a docx from an MBSA report")
parser.add_argument("MBSAfile",help="The name of the .mbsa file")
args=parser.parse_args()

#Globals
MBSAfile = args.MBSAfile
docxname = args.MBSAfile.replace(".mbsa","_report.docx")
xmlroot = ET.parse(MBSAfile).getroot()
domain = xmlroot.get('Domain')
hostname  = xmlroot.get('Machine')
ipaddr = xmlroot.get('IP')
grade = xmlroot.get('Grade')
summarizedInfo = (domain, "\\", hostname, "", ipaddr, "", "Overall Grade: ", grade)
document = Document()
#create doc
def docx():
    document.add_heading('MBSA Report', 1)
    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)
    document.save(docxname)
    readMBSA()
#do some useful stuff
def readMBSA():
        document.add_heading(summarizedInfo, 1)
        paragraph_format = document.styles['Normal'].paragraph_format
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)
        for check in xmlroot.iter('Check'):
            seccheck = check.get('Name')
            seccheck = seccheck.upper()
            paragraph = document.add_paragraph('',style='Heading 1')
            runner=paragraph.add_run('{}'.format(seccheck))
            runner.bold=True
            document.save(docxname)
            for sectionGrade in check.get('Grade'):
                formattedSecGrade = ("Section Grade: " + str(sectionGrade))
                paragraph = document.add_paragraph('',style='Body Text')
                runner = paragraph.add_run('{}'.format(formattedSecGrade))
                runner.bold=True
                table = document.add_table(rows=1, cols=1)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = seccheck
                table.style = 'TableGrid'
                document.save(docxname)
            for advice in check.iter('Advice'):
                easyadvice=(advice.text)
                if easyadvice is "None":
                    easyadvice = "No further information"
                    paragraph = document.add_paragraph('',style='Body Text')
                    runner = paragraph.add_run('{}'.format(easyadvice))
                    runner.bold=False
                    document.save(docxname)
                else:
                    paragraph = document.add_paragraph('',style='Body Text')
                    runner = paragraph.add_run('{}'.format(easyadvice))
                    runner.bold=False
                    document.save(docxname)
            for details in check.iter('Detail'):
                #this part isn't working
                deets = details.get('text')
                if not deets:
                    deets = ""
                    #deets = details.get('text')
                    paragraph = document.add_paragraph('',style='Body Text')
                    runner = paragraph.add_run('{}'.format(deets))
                    runner.bold=False
                    document.save(docxname)
                elif deets is None:
                    deets = ""
                    #deets = details.get('text')
                    paragraph = document.add_paragraph('',style='Body Text')
                    runner = paragraph.add_run('{}'.format(deets))
                    runner.bold=False
                    document.save(docxname)
                else:
                    paragraph = document.add_paragraph('',style='Body Text')
                    runner = paragraph.add_run('{}'.format(deets))
                    runner.bold=False
                    document.save(docxname)
                for rowG in details.iter('Row'):
                    grade = rowG.get('Grade')
                    if grade == 6:
                        grade = "TEST NOT APPLICABLE"
                    for col in rowG.iter('Col'):
                        colText=(col.text)
                        extColText = ("Row Grade: " + grade + "  " + "--" +  "  " + colText)
                        cell = table.add_row().cells
                        cell[0].text = extColText
                        document.save(docxname)
        secUpdates()
def secUpdates():
    #set these globally so we can come back
    run = document.add_paragraph().add_run()
    font = run.font
    font.size = Pt(8)
    for check in xmlroot.iter('Check'):
        seccheck = check.get('Name')
        if "Windows Security Updates" in seccheck:
             for detail in check.iter('Detail'):
                 for data in detail.iter('UpdateData'):
                     #set all relevant data as vars
                     id = data.get('ID')
                     restart = data.get('RestartRequired')
                     severity = data.get('Severity')
                     installed = data.get('IsInstalled')
                     kbid = data.get('KBID')
                     #set severity levels - colors not working
                     if int(severity) == 2:
                         severity = "CRITICAL"
                         font.color.rgb=RGBColor(0xFF,0x00,0xFF)
                     elif int(severity) == 3:
                         severity = "CRITICAL"
                         font.color.rgb=RGBColor(0xFF,0x00,0xFF)
                     elif int(severity) == 4:
                         severity = "Important"
                         font.color.rgb=RGBColor(0xFF,0xC0,0x00)
                     elif int(severity) == 5:
                         severity = "Minor"
                         font.color.rgb=RGBColor(0x00,0x76,0xCF)
                     else:
                         severity = "Very Minor"
                     for title in check.iter('Title'):
                         title = title.text
                     for refs in data.iter('References'):
                         for info in refs.iter('InformationURL'):
                             infoURL = info.text
                         for dl in data.iter('DownloadURL'):
                             dl = dl.text
                     if "false" in installed:
                         secResults = ("Patch: ", title, "NOT INSTALLED", "\n", "This patch is", severity, "\n", "KBID: ", kbid, "\n", "Information URL: ", infoURL, "\n", "Download Link:", dl, "\n")
                         table = document.add_table(rows=1, cols=5)
                         table.alignment = WD_TABLE_ALIGNMENT.RIGHT
                         table.autofit = False
                         table.columns[0].width = Inches(2.5)
                         hdr_cells = table.rows[0].cells
                         hdr_cells[0].text = "Patch"
                         hdr_cells[1].text = "Severity"
                         hdr_cells[2].text = "KBID"
                         hdr_cells[3].text = "Information URL"
                         hdr_cells[4].text = "Download URL"
                         for hdr_cell in table.column_cells(0):
                             hdr_cell.width = Inches(2.5)
                             font.bold = True
                         cell = table.add_row().cells
                         cell[0].text = title
                         cell[1].text = severity
                         cell[2].text = kbid
                         cell[3].text = infoURL
                         cell[4].text = dl
                         table.style = 'TableGrid'
                         document.save(docxname)

                     else:
                        f = open('/dev/null', 'w')
                        sys.stdout = f
docx()
